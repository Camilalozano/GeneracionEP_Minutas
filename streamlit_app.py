import streamlit as st
import pandas as pd
from docx import Document
import zipfile
import io
from datetime import datetime
import time
import base64
from pathlib import Path
import uuid
import hashlib
import os
import re
import unicodedata
from urllib.parse import quote

import requests

# ============== CONFIGURACIÓN DE PÁGINA ==============
st.set_page_config(
    page_title="Atenea: Generador inteligente de Estudios Previos",
    page_icon="Generador de Documentos",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============== ESTILOS CSS PERSONALIZADOS ==============
st.markdown("""
<style>
    /* Fuentes y colores principales */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

    .main-header {
        background: linear-gradient(135deg, #1e3a5f 0%, #2d5a87 100%);
        padding: 2rem;
        border-radius: 16px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 20px rgba(0,0,0,0.1);
    }

    .main-title {
        color: white;
        font-size: 2.5rem;
        font-weight: 700;
        margin: 0;
        display: flex;
        align-items: center;
        gap: 12px;
    }

    .subtitle {
        color: rgba(255,255,255,0.8);
        font-size: 1.1rem;
        margin-top: 0.5rem;
    }

    .stat-card {
        background: linear-gradient(135deg, #f8fafc 0%, #e2e8f0 100%);
        border-radius: 12px;
        padding: 1.5rem;
        text-align: center;
        border-left: 4px solid #2d5a87;
    }

    .step-badge {
        background: #2d5a87;
        color: white;
        padding: 4px 12px;
        border-radius: 20px;
        font-size: 0.8rem;
        font-weight: 600;
    }

    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}

    .sidebar-info {
        background: rgba(255,255,255,0.1);
        border-radius: 8px;
        padding: 1rem;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# ============== FUNCIONES ==============
def replace_text_in_paragraph(paragraph, key, value):
    placeholder = "{{" + key + "}}"
    value_str = str(value) if value is not None else ""
    full_text = "".join(run.text for run in paragraph.runs)
    if placeholder in full_text:
        replaced_text = full_text.replace(placeholder, value_str)
        if paragraph.runs:
            paragraph.runs[0].text = replaced_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(replaced_text)


def diligenciar_documento(word_file, row):
    """Diligencia una plantilla Word con los valores de una fila de datos."""
    word_file.seek(0)
    doc = Document(word_file)

    for paragraph in doc.paragraphs:
        for key in row.index:
            replace_text_in_paragraph(paragraph, key, row[key])

    for table in doc.tables:
        for row_table in table.rows:
            for cell in row_table.cells:
                for paragraph in cell.paragraphs:
                    for key in row.index:
                        replace_text_in_paragraph(paragraph, key, row[key])

    return doc


def guardar_documento_en_bytes(doc):
    """Convierte un documento python-docx en bytes descargables."""
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes


def extraer_texto_documento(doc):
    """Genera una vista previa textual del contenido diligenciado de un documento."""
    bloques = []

    for paragraph in doc.paragraphs:
        texto = paragraph.text.strip()
        if texto:
            bloques.append(texto)

    for table in doc.tables:
        for row_table in table.rows:
            celdas = [cell.text.strip() for cell in row_table.cells if cell.text.strip()]
            if celdas:
                bloques.append(" | ".join(celdas))

    return "\n\n".join(bloques)


def generar_vista_previa_documento(df, word_file, indice=0):
    """Diligencia una sola fila para mostrarla y permitir su descarga en Streamlit."""
    if df is None or df.empty:
        raise ValueError("No hay datos disponibles para generar la vista previa.")

    row = df.iloc[indice]
    doc = diligenciar_documento(word_file, row)
    return doc, guardar_documento_en_bytes(doc)


def generar_documentos(df, word_file, progress_bar, status_text, actor=None, guardar_en_supabase=True):
    try:
        zip_buffer = io.BytesIO()
        documentos_generados = 0
        documentos_guardados_supabase = 0
        errores = []
        total = len(df)

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
            for idx, row in df.iterrows():
                try:
                    doc = diligenciar_documento(word_file, row)
                    doc_bytes = guardar_documento_en_bytes(doc)
                    contenido_docx = doc_bytes.getvalue()
                    nombre_archivo = f"Documento_{idx + 1}.docx"
                    zipf.writestr(nombre_archivo, contenido_docx)
                    documentos_generados += 1

                    if guardar_en_supabase and supabase_storage_configurado():
                        try:
                            metadata = guardar_documento_generado_supabase(
                                row,
                                idx,
                                nombre_archivo,
                                contenido_docx,
                                actor,
                            )
                            if metadata:
                                documentos_guardados_supabase += 1
                        except Exception as error_supabase:
                            errores.append(
                                f"Fila {idx + 1}: documento generado, "
                                f"pero no se guardó en Supabase Storage: {error_supabase}"
                            )

                    progress = (idx + 1) / total
                    progress_bar.progress(progress)
                    status_text.text(f"📄 Procesando documento {idx + 1} de {total}...")

                except Exception as e:
                    errores.append(f"Fila {idx + 1}: {str(e)}")

        status_text.text("✅ ¡Proceso completado!")
        return zip_buffer, documentos_generados, errores, documentos_guardados_supabase

    except Exception as e:
        raise Exception(f"Error procesando documentos: {str(e)}")


PLANTILLA_VARIABLES = [
    "hoy",
    "contratista",
    "descripcion_necesidad",
    "numeroproyecto",
    "nombreproyectoinversion",
    "codigo_objeto",
    "objeto_contrato",
    "tipo_contrato",
    "codigo_unspsc",
    "segmento_unspsc",
    "familia_unspsc",
    "clase_unspsc",
    "plazo_ejecucion_letras",
    "plazo_ejecucion_numero",
    "unidad_plazo",
    "valor_contrato_letras",
    "valor_contrato_numeros",
    "forma_pago",
    "requiere_licencias_si",
    "requiere_licencias_no",
    "obligaciones_especificas",
    "idoneidad",
    "equivalencias",
    "supervisor_cargo",
    "nombre_firma_autorizada",
    "cargo_firma_autorizada",
    "elaboro",
    "reviso",
    "aprobo",
]

CAMPOS_OBLIGATORIOS = [
    "hoy",
    "contratista",
    "descripcion_necesidad",
    "numeroproyecto",
    "nombreproyectoinversion",
    "codigo_objeto",
    "objeto_contrato",
    "tipo_contrato",
    "codigo_unspsc",
    "plazo_ejecucion_letras",
    "plazo_ejecucion_numero",
    "unidad_plazo",
    "valor_contrato_letras",
    "valor_contrato_numeros",
    "forma_pago",
    "obligaciones_especificas",
    "idoneidad",
    "supervisor_cargo",
    "nombre_firma_autorizada",
    "cargo_firma_autorizada",
    "elaboro",
    "reviso",
    "aprobo",
]

ETIQUETAS_CAMPOS = {
    "hoy": "Fecha",
    "contratista": "Contratista",
    "descripcion_necesidad": "Descripción de la necesidad",
    "numeroproyecto": "Número de proyecto",
    "nombreproyectoinversion": "Nombre del proyecto de inversión",
    "codigo_objeto": "Código de objeto",
    "objeto_contrato": "Objeto del contrato",
    "tipo_contrato": "Tipo de contrato",
    "codigo_unspsc": "Código UNSPSC",
    "segmento_unspsc": "Segmento UNSPSC",
    "familia_unspsc": "Familia UNSPSC",
    "clase_unspsc": "Clase UNSPSC",
    "plazo_ejecucion_letras": "Plazo de ejecución en letras",
    "plazo_ejecucion_numero": "Plazo de ejecución en número",
    "unidad_plazo": "Unidad del plazo",
    "valor_contrato_letras": "Valor del contrato en letras",
    "valor_contrato_numeros": "Valor del contrato en números",
    "forma_pago": "Forma de pago",
    "requiere_licencias_si": "Requiere licencias: sí",
    "requiere_licencias_no": "Requiere licencias: no",
    "obligaciones_especificas": "Obligaciones específicas",
    "idoneidad": "Idoneidad",
    "equivalencias": "Equivalencias",
    "supervisor_cargo": "Cargo del supervisor",
    "nombre_firma_autorizada": "Nombre de la firma autorizada",
    "cargo_firma_autorizada": "Cargo de la firma autorizada",
    "elaboro": "Elaboró",
    "reviso": "Revisó",
    "aprobo": "Aprobó",
}

def validar_formulario(data):
    errores = {}

    for campo in CAMPOS_OBLIGATORIOS:
        if not str(data.get(campo, "")).strip():
            errores[campo] = f"{ETIQUETAS_CAMPOS[campo]} es obligatorio."

    if data.get("objeto_contrato") and len(data["objeto_contrato"].strip()) < 20:
        errores["objeto_contrato"] = "El objeto del contrato debe tener al menos 20 caracteres."

    licencias_si = str(data.get("requiere_licencias_si", "")).strip()
    licencias_no = str(data.get("requiere_licencias_no", "")).strip()
    if bool(licencias_si) == bool(licencias_no):
        errores["requiere_licencias"] = "Selecciona solo una opción para indicar si requiere licencias."

    return errores


SUPABASE_TABLE_ESTUDIOS = "EstudiosPrevios"
SUPABASE_TABLE_BITACORA = "BitacoraAuditoria"
SUPABASE_TABLE_DOCUMENTOS = "DocumentosGenerados"
SUPABASE_BUCKET_DOCUMENTOS_DEFAULT = "documentos-generados"
MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
SUPABASE_COLUMNAS_ESTUDIOS = ["id", *PLANTILLA_VARIABLES]
SUPABASE_COLUMNAS_BITACORA = [
    "id_evento",
    "ID_Caso",
    "fecha_hora_utc",
    "actor",
    "accion",
    "detalle",
]
SUPABASE_COLUMNAS_ENTERAS = {"id", "numeroproyecto"}
SUPABASE_COLUMNAS_FECHA = {"hoy", "fecha_hora_utc"}


def obtener_valor_configuracion(nombre):
    """Obtiene un valor desde st.secrets o variables de entorno sin romper la app."""
    try:
        valor = st.secrets.get(nombre, "")
    except Exception:
        valor = ""

    return valor or os.getenv(nombre, "")


def obtener_configuracion_supabase():
    """Lee la configuración de Supabase desde secretos o variables de entorno."""
    url = obtener_valor_configuracion("SUPABASE_URL")
    key = obtener_valor_configuracion("SUPABASE_PUBLISHABLE_KEY")
    admin_password = obtener_valor_configuracion("ADMIN_PASSWORD")

    return str(url).strip().rstrip("/"), str(key).strip(), str(admin_password).strip()


def supabase_configurado():
    url, key, _ = obtener_configuracion_supabase()
    return bool(url and key)


def obtener_bucket_documentos_supabase():
    """Obtiene el bucket de Storage para documentos generados."""
    bucket = obtener_valor_configuracion("SUPABASE_BUCKET_DOCUMENTOS")
    return str(bucket or SUPABASE_BUCKET_DOCUMENTOS_DEFAULT).strip()


def obtener_clave_storage_supabase():
    """Obtiene la service role key requerida para escribir en un bucket privado."""
    return str(obtener_valor_configuracion("SUPABASE_SERVICE_ROLE_KEY")).strip()


def supabase_storage_configurado():
    url, _, _ = obtener_configuracion_supabase()
    return bool(url and obtener_clave_storage_supabase() and obtener_bucket_documentos_supabase())


def limpiar_segmento_ruta(valor, fallback="sin-valor"):
    """Normaliza texto para usarlo como segmento seguro en rutas de Storage."""
    texto = str(valor or "").strip()
    if not texto:
        texto = fallback

    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(caracter for caracter in texto if not unicodedata.combining(caracter))
    texto = re.sub(r"[^A-Za-z0-9._-]+", "-", texto).strip("-._")
    return texto[:80] or fallback


def construir_ruta_documento_storage(row, idx, nombre_archivo):
    """Construye una ruta única y legible para el documento en Supabase Storage."""
    fecha = datetime.utcnow().strftime("%Y/%m/%d")
    id_caso = limpiar_segmento_ruta(row.get("codigo_objeto", "SIN-CODIGO-OBJETO"), "SIN-CODIGO-OBJETO")
    nombre_seguro = limpiar_segmento_ruta(nombre_archivo, f"Documento_{idx + 1}.docx")
    sufijo = datetime.utcnow().strftime("%H%M%S")
    identificador = uuid.uuid4().hex[:8]
    return f"estudios-previos/{fecha}/{id_caso}/fila-{idx + 1}-{sufijo}-{identificador}-{nombre_seguro}"


def subir_documento_storage(ruta_storage, contenido_docx):
    """Sube un archivo .docx al bucket privado de Supabase Storage."""
    url, _, _ = obtener_configuracion_supabase()
    key = obtener_clave_storage_supabase()
    bucket = obtener_bucket_documentos_supabase()

    if not url or not key or not bucket:
        raise ValueError(
            "Faltan SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY/SUPABASE_PUBLISHABLE_KEY "
            "o SUPABASE_BUCKET_DOCUMENTOS para guardar archivos en Storage."
        )

    endpoint = f"{url}/storage/v1/object/{quote(bucket, safe='')}/{quote(ruta_storage, safe='/')}"
    headers = {
        "apikey": key,
        "Authorization": f"Bearer {key}",
        "Content-Type": MIME_DOCX,
        "x-upsert": "false",
    }
    response = requests.post(endpoint, headers=headers, data=contenido_docx, timeout=45)
    response.raise_for_status()
    return response


def guardar_metadata_documento_supabase(metadata):
    """Guarda la referencia del documento en la tabla DocumentosGenerados."""
    supabase_request("POST", SUPABASE_TABLE_DOCUMENTOS, [metadata])


def guardar_documento_generado_supabase(row, idx, nombre_archivo, contenido_docx, actor=None):
    """Sube el .docx a Storage y registra sus metadatos en Supabase."""
    ruta_storage = construir_ruta_documento_storage(row, idx, nombre_archivo)
    subir_documento_storage(ruta_storage, contenido_docx)

    metadata = {
        # La tabla creada tiene id_caso como identity primary key; se omite para que Supabase lo genere.
        "storage_bucket": obtener_bucket_documentos_supabase(),
        "storage_path": ruta_storage,
        "nombre_archivo": nombre_archivo,
        "mime_type": MIME_DOCX,
        "size_bytes": str(len(contenido_docx)),
        "sha256": hashlib.sha256(contenido_docx).hexdigest(),
        "actor": actor if actor else "No especificado",
        "fecha_hora_utc": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"),
    }
    guardar_metadata_documento_supabase(metadata)
    return metadata


def normalizar_valor_supabase(columna, valor):
    """Convierte valores de formulario/Excel a formatos compatibles con Supabase."""
    if pd.isna(valor):
        return None

    if hasattr(valor, "isoformat") and columna in SUPABASE_COLUMNAS_FECHA:
        return valor.isoformat()

    texto = str(valor).strip()
    if not texto:
        return None

    if columna in SUPABASE_COLUMNAS_ENTERAS:
        try:
            return int(float(texto.replace(",", "")))
        except ValueError:
            return None

    if columna in SUPABASE_COLUMNAS_FECHA:
        fecha = pd.to_datetime(texto, errors="coerce")
        if pd.notna(fecha):
            return fecha.isoformat()

    return texto


def normalizar_registro_estudios(origen):
    """Prepara un registro para insertarlo en la tabla EstudiosPrevios."""
    registro = {}
    for columna in SUPABASE_COLUMNAS_ESTUDIOS:
        if columna == "id":
            continue
        registro[columna] = normalizar_valor_supabase(columna, origen.get(columna, ""))
    return registro


def supabase_request(metodo, tabla, payload=None, params=None):
    """Ejecuta una petición REST contra Supabase."""
    url, publishable_key, _ = obtener_configuracion_supabase()
    key = obtener_clave_storage_supabase() or publishable_key
    if not url or not key:
        raise ValueError(
            "Faltan SUPABASE_URL y SUPABASE_SERVICE_ROLE_KEY/SUPABASE_PUBLISHABLE_KEY "
            "en los secretos de Streamlit."
        )

    endpoint = f"{url}/rest/v1/{tabla}"
    headers = {
        "apikey": key,
        "Authorization": f"Bearer {key}",
        "Content-Type": "application/json",
        "Accept": "application/json",
    }
    if metodo.upper() == "POST":
        headers["Prefer"] = "return=minimal"

    response = requests.request(
        metodo,
        endpoint,
        headers=headers,
        json=payload,
        params=params,
        timeout=20,
    )
    response.raise_for_status()
    return response


def guardar_estudios_previos_supabase(registros, actor=None):
    """Guarda uno o varios registros capturados en EstudiosPrevios."""
    if not registros:
        return 0

    payload = [normalizar_registro_estudios(registro) for registro in registros]
    supabase_request("POST", SUPABASE_TABLE_ESTUDIOS, payload)
    return len(payload)


def guardar_bitacora_supabase(evento):
    """Guarda un evento de auditoría en BitacoraAuditoria."""
    payload = {
        columna: normalizar_valor_supabase(columna, evento.get(columna, ""))
        for columna in SUPABASE_COLUMNAS_BITACORA
        if columna != "id_evento"
    }
    supabase_request("POST", SUPABASE_TABLE_BITACORA, payload)


def consultar_tabla_supabase(tabla, columnas="*", orden=None):
    """Consulta una tabla completa desde Supabase para reportes administrativos."""
    params = {"select": columnas}
    if orden:
        params["order"] = orden

    response = supabase_request("GET", tabla, params=params)
    return pd.DataFrame(response.json())


def construir_dataframe_desde_formulario(data):
    fila = {campo: str(data.get(campo, "")).strip() for campo in PLANTILLA_VARIABLES}
    return pd.DataFrame([fila], columns=PLANTILLA_VARIABLES)


def obtener_fecha_borrador(defaults):
    fecha_guardada = defaults.get("hoy")
    if hasattr(fecha_guardada, "strftime"):
        return fecha_guardada
    if isinstance(fecha_guardada, str) and fecha_guardada.strip():
        for formato in ("%d/%m/%Y", "%Y-%m-%d"):
            try:
                return datetime.strptime(fecha_guardada.strip(), formato).date()
            except ValueError:
                continue
    return datetime.now().date()


if "form_borrador" not in st.session_state:
    st.session_state.form_borrador = {}
if "df_captura" not in st.session_state:
    st.session_state.df_captura = None

if "resultado_zip" not in st.session_state:
    st.session_state.resultado_zip = None
if "resultado_nombre" not in st.session_state:
    st.session_state.resultado_nombre = ""
if "resultado_generados" not in st.session_state:
    st.session_state.resultado_generados = 0
if "resultado_errores" not in st.session_state:
    st.session_state.resultado_errores = []
if "resultado_guardados_supabase" not in st.session_state:
    st.session_state.resultado_guardados_supabase = 0
if "descarga_automatica_pendiente" not in st.session_state:
    st.session_state.descarga_automatica_pendiente = False
if "auditoria_acciones" not in st.session_state:
    st.session_state.auditoria_acciones = []
if "excel_guardado_hash" not in st.session_state:
    st.session_state.excel_guardado_hash = ""


def obtener_id_caso_desde_codigo_objeto(origen=None):
    """Obtiene el ID del caso desde la variable de plantilla {{codigo_objeto}}."""
    if isinstance(origen, pd.DataFrame) and "codigo_objeto" in origen.columns and not origen.empty:
        valor = origen.iloc[0].get("codigo_objeto", "")
    elif isinstance(origen, dict):
        valor = origen.get("codigo_objeto", "")
    else:
        valor = st.session_state.form_borrador.get("codigo_objeto", "")

    valor_normalizado = str(valor).strip()
    return valor_normalizado if valor_normalizado else "SIN-CODIGO-OBJETO"


def registrar_evento_auditoria(accion, actor, detalle, id_caso=None):
    evento = {
        "id_evento": str(uuid.uuid4())[:8],
        "ID_Caso": id_caso or obtener_id_caso_desde_codigo_objeto(),
        "fecha_hora_utc": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"),
        "actor": actor if actor else "No especificado",
        "accion": accion,
        "detalle": detalle,
    }
    st.session_state.auditoria_acciones.append(evento)

    if supabase_configurado():
        try:
            guardar_bitacora_supabase(evento)
        except Exception as error:
            st.warning(f"No fue posible guardar la bitácora en Supabase: {error}")


def cargar_plantilla_precargada():
    """Carga la plantilla institucional por defecto desde el repositorio."""
    candidatos = [
        Path("Plantilla.docx"),
        Path("templates/Plantilla.docx"),
    ]

    for ruta in candidatos:
        if ruta.exists() and ruta.is_file():
            return io.BytesIO(ruta.read_bytes()), ruta.name

    return None, None


def disparar_descarga_automatica(zip_bytes, file_name):
    b64 = base64.b64encode(zip_bytes).decode()
    st.markdown(
        f"""
        <a id="auto-download-link" href="data:application/zip;base64,{b64}" download="{file_name}"></a>
        <script>
            const enlace = document.getElementById('auto-download-link');
            if (enlace) {{ enlace.click(); }}
        </script>
        """,
        unsafe_allow_html=True,
    )

# ============== SIDEBAR ==============
with st.sidebar:
    st.markdown("## Generador Inteligente de Documentos")
    st.markdown("---")
    st.markdown("### 📖 Guía Rápida")
    st.markdown("### ℹ️ Información")
    st.markdown("""
    <div class="sidebar-info">
        <p><strong>Versión:</strong> 1.1</p>
        <p><strong>Gerencia de Gestión Corporativa</p>
    </div>
    """, unsafe_allow_html=True)

    if supabase_configurado():
        st.success("✅ Supabase configurado")
        if supabase_storage_configurado():
            st.success(f"🗂️ Storage: {obtener_bucket_documentos_supabase()}")
        else:
            st.warning("⚠️ Storage no configurado para guardar .docx. Agrega SUPABASE_SERVICE_ROLE_KEY.")
    else:
        st.warning("⚠️ Supabase no configurado. Revisa los secretos de Streamlit.")

# Header
st.markdown("""
<div class="main-header">
    <h1 class="main-title" style="color: white;">Atenea: Generador inteligente de Estudios Previos</h1>
    <p class="subtitle">Estudios Previos | Gerencia de Gestión Corporativa</p>
</div>
""", unsafe_allow_html=True)

actor_actual = st.text_input(
    "👤 Responsable de la acción (usuario actual)",
    value=st.session_state.get("actor_actual", ""),
    help="Este nombre se registra en la bitácora para control institucional y trazabilidad legal.",
)
st.session_state.actor_actual = actor_actual


st.markdown("<br>", unsafe_allow_html=True)
st.markdown("### 📁 Captura de datos")
modo_captura = st.radio(
    "Selecciona el modo de captura",
    ["Formulario guiado (principal)", "Carga masiva por Excel (secundario)"],
    horizontal=True,
)

excel_file = None
df = None

if modo_captura == "Formulario guiado (principal)":
    st.info(
        "🧭 Diligencia los campos del formulario guiado. "
        "Cada campo corresponde a una variable disponible en la plantilla Word."
    )

    defaults = st.session_state.form_borrador
    with st.form("formulario_guiado"):
        st.markdown("#### 1) Identificación del proceso")
        id_col1, id_col2 = st.columns(2)
        with id_col1:
            hoy = st.date_input(
                "Fecha *",
                value=obtener_fecha_borrador(defaults),
                help="Variable: {{hoy}}",
            )
            contratista = st.text_input(
                "Contratista *",
                value=defaults.get("contratista", ""),
                help="Variable: {{contratista}}",
            )
            numeroproyecto = st.text_input(
                "Número de proyecto *",
                value=defaults.get("numeroproyecto", ""),
                help="Variable: {{numeroproyecto}}",
            )
        with id_col2:
            nombreproyectoinversion = st.text_input(
                "Nombre del proyecto de inversión *",
                value=defaults.get("nombreproyectoinversion", ""),
                help="Variable: {{nombreproyectoinversion}}",
            )
            codigo_objeto = st.text_input(
                "Código de objeto *",
                value=defaults.get("codigo_objeto", ""),
                help="Variable: {{codigo_objeto}}",
            )
            tipo_contrato = st.text_input(
                "Tipo de contrato *",
                value=defaults.get("tipo_contrato", ""),
                help="Variable: {{tipo_contrato}}",
            )

        st.markdown("#### 2) Necesidad y objeto contractual")
        descripcion_necesidad = st.text_area(
            "Descripción de la necesidad *",
            value=defaults.get("descripcion_necesidad", ""),
            help="Variable: {{descripcion_necesidad}}",
        )
        objeto_contrato = st.text_area(
            "Objeto del contrato *",
            value=defaults.get("objeto_contrato", ""),
            help="Variable: {{objeto_contrato}}",
        )

        st.markdown("#### 3) Clasificación UNSPSC")
        unspsc_col1, unspsc_col2, unspsc_col3, unspsc_col4 = st.columns(4)
        with unspsc_col1:
            codigo_unspsc = st.text_input(
                "Código UNSPSC *",
                value=defaults.get("codigo_unspsc", ""),
                help="Variable: {{codigo_unspsc}}",
            )
        with unspsc_col2:
            segmento_unspsc = st.text_input(
                "Segmento UNSPSC",
                value=defaults.get("segmento_unspsc", ""),
                help="Variable: {{segmento_unspsc}}",
            )
        with unspsc_col3:
            familia_unspsc = st.text_input(
                "Familia UNSPSC",
                value=defaults.get("familia_unspsc", ""),
                help="Variable: {{familia_unspsc}}",
            )
        with unspsc_col4:
            clase_unspsc = st.text_input(
                "Clase UNSPSC",
                value=defaults.get("clase_unspsc", ""),
                help="Variable: {{clase_unspsc}}",
            )

        st.markdown("#### 4) Plazo, valor y pago")
        plazo_col1, plazo_col2, plazo_col3 = st.columns(3)
        with plazo_col1:
            plazo_ejecucion_letras = st.text_input(
                "Plazo de ejecución en letras *",
                value=defaults.get("plazo_ejecucion_letras", ""),
                help="Variable: {{plazo_ejecucion_letras}}",
            )
        with plazo_col2:
            plazo_ejecucion_numero = st.text_input(
                "Plazo de ejecución en número *",
                value=defaults.get("plazo_ejecucion_numero", ""),
                help="Variable: {{plazo_ejecucion_numero}}",
            )
        with plazo_col3:
            unidad_plazo = st.text_input(
                "Unidad del plazo *",
                value=defaults.get("unidad_plazo", ""),
                help="Variable: {{unidad_plazo}}",
            )

        valor_col1, valor_col2 = st.columns(2)
        with valor_col1:
            valor_contrato_letras = st.text_input(
                "Valor del contrato en letras *",
                value=defaults.get("valor_contrato_letras", ""),
                help="Variable: {{valor_contrato_letras}}",
            )
        with valor_col2:
            valor_contrato_numeros = st.text_input(
                "Valor del contrato en números *",
                value=defaults.get("valor_contrato_numeros", ""),
                help="Variable: {{valor_contrato_numeros}}",
            )
        forma_pago = st.text_area(
            "Forma de pago *",
            value=defaults.get("forma_pago", ""),
            help="Variable: {{forma_pago}}",
        )

        st.markdown("#### 5) Licencias y obligaciones")
        requiere_licencias = st.radio(
            "¿Requiere licencias? *",
            ["Sí", "No"],
            index=0 if defaults.get("requiere_licencias_si") else 1,
            horizontal=True,
            help="Marca una opción para diligenciar {{requiere_licencias_si}} o {{requiere_licencias_no}} con una X.",
        )
        obligaciones_especificas = st.text_area(
            "Obligaciones específicas *",
            value=defaults.get("obligaciones_especificas", ""),
            help="Variable: {{obligaciones_especificas}}",
            height=180,
        )

        st.markdown("#### 6) Idoneidad, supervisión y firmas")
        idoneidad = st.text_area(
            "Idoneidad *",
            value=defaults.get("idoneidad", ""),
            help="Variable: {{idoneidad}}",
        )
        equivalencias = st.text_area(
            "Equivalencias",
            value=defaults.get("equivalencias", ""),
            help="Variable: {{equivalencias}}",
        )
        supervisor_cargo = st.text_input(
            "Cargo del supervisor *",
            value=defaults.get("supervisor_cargo", ""),
            help="Variable: {{supervisor_cargo}}",
        )

        firma_col1, firma_col2 = st.columns(2)
        with firma_col1:
            nombre_firma_autorizada = st.text_input(
                "Nombre de la firma autorizada *",
                value=defaults.get("nombre_firma_autorizada", ""),
                help="Variable: {{nombre_firma_autorizada}}",
            )
        with firma_col2:
            cargo_firma_autorizada = st.text_input(
                "Cargo de la firma autorizada *",
                value=defaults.get("cargo_firma_autorizada", ""),
                help="Variable: {{cargo_firma_autorizada}}",
            )

        st.markdown("#### 7) Control de elaboración")
        ctrl_col1, ctrl_col2, ctrl_col3 = st.columns(3)
        with ctrl_col1:
            elaboro = st.text_input(
                "Elaboró *",
                value=defaults.get("elaboro", ""),
                help="Variable: {{elaboro}}",
            )
        with ctrl_col2:
            reviso = st.text_input(
                "Revisó *",
                value=defaults.get("reviso", ""),
                help="Variable: {{reviso}}",
            )
        with ctrl_col3:
            aprobo = st.text_input(
                "Aprobó *",
                value=defaults.get("aprobo", ""),
                help="Variable: {{aprobo}}",
            )

        c1, c2 = st.columns(2)
        guardar_borrador = c1.form_submit_button("💾 Guardar borrador")
        cargar_registro = c2.form_submit_button("✅ Usar este registro")

    form_data = {
        "hoy": hoy.strftime("%d/%m/%Y"),
        "contratista": contratista,
        "descripcion_necesidad": descripcion_necesidad,
        "numeroproyecto": numeroproyecto,
        "nombreproyectoinversion": nombreproyectoinversion,
        "codigo_objeto": codigo_objeto,
        "objeto_contrato": objeto_contrato,
        "tipo_contrato": tipo_contrato,
        "codigo_unspsc": codigo_unspsc,
        "segmento_unspsc": segmento_unspsc,
        "familia_unspsc": familia_unspsc,
        "clase_unspsc": clase_unspsc,
        "plazo_ejecucion_letras": plazo_ejecucion_letras,
        "plazo_ejecucion_numero": plazo_ejecucion_numero,
        "unidad_plazo": unidad_plazo,
        "valor_contrato_letras": valor_contrato_letras,
        "valor_contrato_numeros": valor_contrato_numeros,
        "forma_pago": forma_pago,
        "requiere_licencias_si": "X" if requiere_licencias == "Sí" else "",
        "requiere_licencias_no": "X" if requiere_licencias == "No" else "",
        "obligaciones_especificas": obligaciones_especificas,
        "idoneidad": idoneidad,
        "equivalencias": equivalencias,
        "supervisor_cargo": supervisor_cargo,
        "nombre_firma_autorizada": nombre_firma_autorizada,
        "cargo_firma_autorizada": cargo_firma_autorizada,
        "elaboro": elaboro,
        "reviso": reviso,
        "aprobo": aprobo,
    }
    errores = validar_formulario(form_data)

    if errores and cargar_registro:
        with st.expander("⚠️ Validaciones pendientes", expanded=True):
            for mensaje in errores.values():
                st.error(mensaje)

    if guardar_borrador:
        st.session_state.form_borrador = form_data
        registrar_evento_auditoria(
            "Guardar borrador",
            actor_actual,
            "Se guardó el borrador del formulario guiado con variables de la plantilla.",
            obtener_id_caso_desde_codigo_objeto(form_data),
        )
        st.success("✅ Borrador guardado en la sesión actual.")

    if cargar_registro:
        if errores:
            st.warning("⚠️ Corrige las validaciones antes de usar el registro.")
        else:
            st.session_state.df_captura = construir_dataframe_desde_formulario(form_data)
            if supabase_configurado():
                try:
                    guardar_estudios_previos_supabase([form_data], actor_actual)
                    st.info("💾 Registro guardado en Supabase (EstudiosPrevios).")
                except Exception as error:
                    st.error(f"No fue posible guardar el registro en Supabase: {error}")
            registrar_evento_auditoria(
                "Cargar registro",
                actor_actual,
                "Se cargó un registro del formulario guiado para generar documentos.",
                obtener_id_caso_desde_codigo_objeto(form_data),
            )
            st.success("✅ Registro cargado correctamente para generar documentos.")

    if st.session_state.df_captura is not None:
        df = st.session_state.df_captura

        with st.expander("👀 Vista previa del registro capturado", expanded=True):
            st.dataframe(st.session_state.df_captura, use_container_width=True)

else:
    st.markdown("##### 📊 Datos (Excel)")
    excel_file = st.file_uploader(
        "Arrastra o selecciona tu archivo Excel",
        type="xlsx",
        key="excel",
        help="Archivo Excel con múltiples registros para generación masiva"
    )
    if excel_file:
        try:
            excel_bytes = excel_file.getvalue()
            excel_hash = hashlib.sha256(excel_bytes).hexdigest()
            df = pd.read_excel(io.BytesIO(excel_bytes))

            if supabase_configurado() and st.session_state.excel_guardado_hash != excel_hash:
                registros_excel = df.fillna("").to_dict(orient="records")
                guardados = guardar_estudios_previos_supabase(registros_excel, actor_actual)
                st.session_state.excel_guardado_hash = excel_hash
                st.info(f"💾 {guardados} registros guardados en Supabase (EstudiosPrevios).")

            registrar_evento_auditoria(
                "Cargar Excel",
                actor_actual,
                f"Se cargó archivo Excel: {excel_file.name} con {len(df)} registros.",
                obtener_id_caso_desde_codigo_objeto(df),
            )
            st.success(f"✅ {excel_file.name}")
        except Exception as e:
            st.error(f"Error al leer Excel: {e}")

st.markdown("##### 📝 Plantilla (Word)")
plantilla_precargada, nombre_plantilla_precargada = cargar_plantilla_precargada()

if plantilla_precargada:
    st.success(f"✅ Plantilla precargada disponible: {nombre_plantilla_precargada}")
    st.info("ℹ️ Se usará siempre la plantilla precargada por defecto.")
else:
    st.warning("⚠️ No se encontró la plantilla precargada Plantilla.docx.")

word_file = plantilla_precargada

if df is not None and word_file:
    with st.expander("👀 Vista previa de datos", expanded=True):
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("📄 Documentos a generar", len(df))
        with col2:
            st.metric("📊 Columnas/Variables", len(df.columns))
        with col3:
            st.metric("🔤 Placeholders detectados", len(df.columns))
        st.dataframe(df, use_container_width=True, height=220)
        placeholders = " | ".join([f"`{{{{{col}}}}}`" for col in df.columns])
        st.code(placeholders, language=None)

    st.markdown("### 👁️ Vista previa de la plantilla diligenciada")
    st.caption(
        "Puedes revisar cómo quedará la plantilla con los datos del primer registro "
        "antes de generar el paquete completo."
    )

    preview_col1, preview_col2 = st.columns([1, 2])
    with preview_col1:
        registro_preview = st.number_input(
            "Registro para previsualizar",
            min_value=1,
            max_value=len(df),
            value=1,
            step=1,
            help="Selecciona la fila de datos que quieres usar para diligenciar la vista previa.",
        )

    try:
        doc_preview, doc_preview_bytes = generar_vista_previa_documento(
            df,
            word_file,
            int(registro_preview) - 1,
        )
        texto_preview = extraer_texto_documento(doc_preview)

        with st.expander("📄 Ver contenido diligenciado", expanded=True):
            if texto_preview:
                st.text_area(
                    "Contenido extraído de la plantilla diligenciada",
                    value=texto_preview,
                    height=360,
                    disabled=True,
                    help="Vista textual de revisión. El archivo descargable conserva el formato Word de la plantilla.",
                )
            else:
                st.info(
                    "La plantilla se diligenció, pero no se detectó texto extraíble para mostrar. "
                    "Descarga la vista previa en Word para revisarla con formato."
                )

        with preview_col2:
            st.download_button(
                label="📥 Descargar vista previa diligenciada (.docx)",
                data=doc_preview_bytes.getvalue(),
                file_name=f"Vista_previa_documento_{int(registro_preview)}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
    except Exception as e:
        st.warning(f"No fue posible generar la vista previa diligenciada: {e}")

    generate_btn = st.button("🚀 Generar Documentos", use_container_width=True, type="primary")
    if generate_btn:
        if df.empty:
            st.error("❌ No hay datos para procesar")
        else:
            progress_bar = st.progress(0)
            status_text = st.empty()
            zip_buffer, generados, errores, guardados_supabase = generar_documentos(
                df,
                word_file,
                progress_bar,
                status_text,
                actor_actual,
            )
            time.sleep(0.5)
            if generados > 0:
                zip_buffer.seek(0)
                zip_data = zip_buffer.getvalue()
                nombre_zip = f"Documentos_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"

                st.session_state.resultado_zip = zip_data
                st.session_state.resultado_nombre = nombre_zip
                st.session_state.resultado_generados = generados
                st.session_state.resultado_errores = errores
                st.session_state.resultado_guardados_supabase = guardados_supabase
                st.session_state.descarga_automatica_pendiente = True
                registrar_evento_auditoria(
                    "Generar documentos",
                    actor_actual,
                    f"Se generaron {generados} documentos, "
                    f"se guardaron {guardados_supabase} en Supabase Storage "
                    f"y se registraron {len(errores)} errores.",
                    obtener_id_caso_desde_codigo_objeto(df),
                )

                if supabase_storage_configurado():
                    st.info(f"🗂️ {guardados_supabase} documento(s) guardado(s) en Supabase Storage.")
                st.success("✅ Documentos generados correctamente. Iniciando descarga automática...")
            else:
                st.error("❌ No se pudieron generar documentos")

if st.session_state.resultado_zip:
    st.download_button(
        label=(
            f"📥 Descargar resultados ({st.session_state.resultado_generados} documentos; "
            f"{st.session_state.resultado_guardados_supabase} guardados en Supabase)"
        ),
        data=st.session_state.resultado_zip,
        file_name=st.session_state.resultado_nombre,
        mime="application/zip",
        use_container_width=True,
        type="secondary"
    )

    if st.session_state.descarga_automatica_pendiente:
        disparar_descarga_automatica(st.session_state.resultado_zip, st.session_state.resultado_nombre)
        registrar_evento_auditoria(
            "Descarga automática",
            actor_actual,
            f"Se disparó descarga automática del archivo {st.session_state.resultado_nombre}.",
            obtener_id_caso_desde_codigo_objeto(df),
        )
        st.session_state.descarga_automatica_pendiente = False

    if st.session_state.resultado_errores:
        with st.expander(f"⚠️ Ver {len(st.session_state.resultado_errores)} errores"):
            for error in st.session_state.resultado_errores:
                st.warning(error)

st.markdown("---")
st.markdown("### 🔐 Panel administrativo de Supabase")
st.caption(
    "Consulta y descarga la información guardada en EstudiosPrevios y BitacoraAuditoria. "
    "Este panel requiere la clave ADMIN_PASSWORD configurada en los secretos de Streamlit."
)

_, _, admin_password_configurada = obtener_configuracion_supabase()
clave_admin = st.text_input("Clave de administrador", type="password")

if not supabase_configurado():
    st.info("Configura SUPABASE_URL y SUPABASE_PUBLISHABLE_KEY para habilitar este panel.")
elif not admin_password_configurada:
    st.warning("Configura ADMIN_PASSWORD en los secretos de Streamlit para habilitar descargas administrativas.")
elif clave_admin == admin_password_configurada:
    tab_estudios, tab_bitacora, tab_documentos = st.tabs([
        "📄 Estudios previos",
        "🧾 Bitácora",
        "🗂️ Documentos generados",
    ])

    with tab_estudios:
        try:
            df_estudios_supabase = consultar_tabla_supabase(SUPABASE_TABLE_ESTUDIOS, orden="id.desc")
            st.dataframe(df_estudios_supabase, use_container_width=True, height=260)
            st.download_button(
                label="📥 Descargar EstudiosPrevios (CSV)",
                data=df_estudios_supabase.to_csv(index=False).encode("utf-8"),
                file_name=f"EstudiosPrevios_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
                use_container_width=True,
            )
        except Exception as error:
            st.error(f"No fue posible consultar EstudiosPrevios en Supabase: {error}")

    with tab_bitacora:
        try:
            df_bitacora_supabase = consultar_tabla_supabase(SUPABASE_TABLE_BITACORA, orden="fecha_hora_utc.desc")
            st.dataframe(df_bitacora_supabase, use_container_width=True, height=260)
            st.download_button(
                label="📥 Descargar BitacoraAuditoria (CSV)",
                data=df_bitacora_supabase.to_csv(index=False).encode("utf-8"),
                file_name=f"BitacoraAuditoria_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
                use_container_width=True,
            )
        except Exception as error:
            st.error(f"No fue posible consultar BitacoraAuditoria en Supabase: {error}")

    with tab_documentos:
        try:
            df_documentos_supabase = consultar_tabla_supabase(SUPABASE_TABLE_DOCUMENTOS, orden="fecha_hora_utc.desc")
            st.dataframe(df_documentos_supabase, use_container_width=True, height=260)
            st.download_button(
                label="📥 Descargar DocumentosGenerados (CSV)",
                data=df_documentos_supabase.to_csv(index=False).encode("utf-8"),
                file_name=f"DocumentosGenerados_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
                use_container_width=True,
            )
        except Exception as error:
            st.error(f"No fue posible consultar DocumentosGenerados en Supabase: {error}")
elif clave_admin:
    st.error("Clave de administrador incorrecta.")
